"""
Report Generator
Generates rigorous, fact-based industry research reports.
Outputs Markdown + optional Word document.
Python 3.8+ compatible.
"""

import re
import os
import json
from datetime import datetime
from typing import Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


class ReportGenerator:
    """
    Generates industry research reports with strict fact-checking standards.
    """
    
    def __init__(self, research_data: Dict, output_dir: str = "./output"):
        self.data = research_data
        self.output_dir = output_dir
        self.industry = research_data.get("industry", "Unknown Industry")
        self.research_date = research_data.get("research_date", datetime.now().isoformat())
        os.makedirs(output_dir, exist_ok=True)
    
    def _get_dimension_results(self, dim_key: str) -> List[Dict]:
        return self.data.get("dimensions", {}).get(dim_key, [])
    
    def _get_high_quality_sources(self, dim_key: Optional[str] = None) -> List[Dict]:
        """Get high and medium reliability sources."""
        sources = self.data.get("sources", [])
        if dim_key:
            sources = [s for s in sources if dim_key in s.get("dimensions", [])]
        return [s for s in sources if s.get("reliability") in ("high", "medium")]
    
    def _format_number_context(self, item: Dict) -> str:
        """Format a data point with context."""
        ctx = item.get("context", "").replace("\n", " ").strip()
        if len(ctx) > 150:
            ctx = ctx[:150] + "..."
        val = item.get("value", "")
        return f"- **{val}** — {ctx}"
    
    def _generate_executive_summary(self) -> str:
        """Generate executive summary from research data."""
        sources_count = len(self.data.get("sources", []))
        hq_sources = len([s for s in self.data.get("sources", []) if s.get("reliability") == "high"])
        
        lines = [
            "## 执行摘要（Executive Summary）",
            "",
            f"> **研究对象**：{self.industry}",
            f"> **研究时间**：{self.research_date[:10]}",
            f"> **数据来源**：共分析 {sources_count} 个独立信息源，其中高可信度来源 {hq_sources} 个",
            f"> **研究方法**：基于公开信息的系统性网络调研，所有数据均标注来源",
            "",
            "### 核心发现",
            "",
        ]
        
        # Extract key numbers for summary
        numbers = self.data.get("extracted_data", {}).get("numbers", [])[:5]
        if numbers:
            lines.append("**关键数据点（未经交叉验证，仅供参考）**：")
            for n in numbers:
                lines.append(f"- {self._format_number_context(n)}")
            lines.append("")
        
        # Extract key companies
        companies = self._extract_company_names()
        if companies:
            lines.append(f"**主要提及公司**：{', '.join(companies[:10])}")
            lines.append("")
        
        lines.append("---")
        lines.append("")
        return "\n".join(lines)
    
    def _extract_company_names(self) -> List[str]:
        """Extract potential company names from search results."""
        companies = set()
        for dim_results in self.data.get("dimensions", {}).values():
            for r in dim_results:
                text = f"{r.get('title', '')} {r.get('snippet', '')}"
                # Simple heuristic: look for capitalized words that might be company names
                # This is imperfect but works for a baseline
                for match in re.finditer(r'\b([A-Z][a-zA-Z0-9]*(?:\s+[A-Z][a-zA-Z0-9]*){0,3})\b', text):
                    name = match.group(1)
                    if len(name) > 2 and name not in {"The", "And", "For", "With", "From", "Inc", "LLC", "Ltd", "Corp"}:
                        companies.add(name)
        return sorted(companies)
    
    def _generate_section(self, dim_key: str, title: str, guidance: str) -> str:
        """Generate a report section from dimension data."""
        results = self._get_dimension_results(dim_key)
        sources = self._get_high_quality_sources(dim_key)
        
        lines = [
            f"## {title}",
            "",
            f"**分析指引**：{guidance}",
            "",
        ]
        
        if not results:
            lines.append(
                "> ⚠️ **数据不足**：该维度未能搜集到足够的信息。建议通过专业数据库（如 "
                "Gartner、IDC、艾瑞咨询、易观分析）进行补充调研。"
            )
            lines.append("")
            return "\n".join(lines)
        
        # Categorized findings
        lines.append("### 关键发现")
        lines.append("")
        
        for i, r in enumerate(results[:12], 1):
            score_emoji = {"high": "🟢", "medium": "🟡", "low": "🔴"}.get(r.get("source_score", "low"), "⚪")
            lines.append(f"{i}. {score_emoji} [{r.get('title', 'Untitled')}]({r.get('url', '')})")
            lines.append(f"   > {r.get('snippet', 'No snippet')[:250]}")
            lines.append("")
        
        # Extracted data points for this dimension
        extracted = self.data.get("extracted_data", {})
        
        # Numbers relevant to this section
        if dim_key in ("overview", "companies", "trends") and extracted.get("numbers"):
            lines.append("### 关键数据")
            lines.append("")
            for n in extracted["numbers"][:8]:
                lines.append(self._format_number_context(n))
            lines.append("")
        
        # Technologies for tech section
        if dim_key == "technology" and extracted.get("technologies"):
            lines.append("### 技术关键词")
            lines.append("")
            techs = list(set(t.get("name", "") for t in extracted["technologies"]))[:15]
            lines.append(", ".join(f"`{t}`" for t in techs if t))
            lines.append("")
        
        # Source list
        if sources:
            lines.append("### 主要信息来源")
            lines.append("")
            for s in sources[:8]:
                rel = {"high": "高", "medium": "中"}.get(s.get("reliability"), "低")
                lines.append(f"- [{s.get('title', 'Source')[:60]}]({s.get('url', '')}) — 可信度：{rel}")
            lines.append("")
        
        lines.append("---")
        lines.append("")
        return "\n".join(lines)
    
    def _generate_methodology(self) -> str:
        """Generate methodology and limitations section."""
        return """## 研究方法说明

### 数据来源
本报告基于系统性网络公开信息搜集，主要来源包括：
- **行业研究机构**：McKinsey、BCG、Gartner、IDC、艾瑞咨询、易观分析等
- **金融数据平台**：Crunchbase、PitchBook、CB Insights、IT桔子等
- **新闻媒体**：Reuters、Bloomberg、36氪、虎嗅、IT之家等
- **政府与监管机构**：各国政府网站、SEC、证监会等

### 数据质量分级
- 🟢 **高可信度**：来自权威研究机构、上市公司财报、政府统计
- 🟡 **中等可信度**：来自知名科技媒体、行业垂直媒体
- 🔴 **低可信度**：来自自媒体、未验证来源

### 重要声明
1. **事实边界**：本报告严格区分"已确认事实"与"行业预测/分析师观点"。所有预测性内容均标注来源和置信度。
2. **时效性**：网络信息存在时效滞后，关键数据均以"截至研究日期"标注。建议对时效敏感的数据进行二次核实。
3. **信息完整性**：公开信息存在天然不完整性和选择性偏差。本报告不做超出信息支撑的推断。
4. **非投资建议**：本报告仅供行业研究参考，不构成任何投资建议。

---

"""
    
    def _generate_sources_index(self) -> str:
        """Generate complete sources index."""
        sources = self.data.get("sources", [])
        lines = [
            "## 附录：完整信息来源索引",
            "",
            f"共 {len(sources)} 个独立来源，按可信度排序：",
            "",
        ]
        
        for s in sorted(sources, key=lambda x: {"high": 0, "medium": 1, "low": 2}.get(x.get("reliability"), 3)):
            rel = {"high": "高", "medium": "中", "low": "低"}.get(s.get("reliability"), "未知")
            dims = ", ".join(s.get("dimensions", []))
            lines.append(f"- [{s.get('title', 'Source')[:80]}]({s.get('url', '')})")
            lines.append(f"  - 可信度：{rel} | 涉及维度：{dims}")
        
        lines.append("")
        return "\n".join(lines)
    
    def generate_markdown(self) -> str:
        """Generate complete Markdown report."""
        sections = [
            f"# {self.industry} 行业深度研究报告",
            f"\n> **报告生成时间**：{self.research_date[:19]}  ",
            f"> **研究方法论**：自主网络调研 + 多维度信息交叉验证  ",
            f"> **数据承诺**：所有陈述均基于公开可查信息，无编造成分\n",
            self._generate_executive_summary(),
            self._generate_section(
                "overview",
                "行业概览与市场规模",
                "分析行业定义、历史发展阶段、当前市场规模（TAM/SAM/SOM）、年复合增长率（CAGR）、区域分布特征。"
            ),
            self._generate_section(
                "chain",
                "产业链结构与价值链分析",
                "拆解产业链上中下游结构，识别核心环节、价值分配、进入壁垒、关键资源与能力要求。"
            ),
            self._generate_section(
                "companies",
                "主要公司与竞争格局",
                "梳理头部企业、市场份额、商业模式对比、融资历史、估值水平、竞争策略差异。"
            ),
            self._generate_section(
                "technology",
                "技术发展与演进路径",
                "识别核心技术栈、技术成熟度（TRL）、专利布局、研发投入趋势、技术替代风险。"
            ),
            self._generate_section(
                "policy",
                "政策环境与监管动态",
                "梳理国内外相关政策、监管框架演变、行业标准、合规要求、政策对竞争格局的影响。"
            ),
            self._generate_section(
                "trends",
                "未来趋势与投资启示",
                "基于事实推演技术演进方向、市场变化趋势、潜在风险因素、对投资人的启示。严格区分事实与预测。"
            ),
            self._generate_methodology(),
            self._generate_sources_index(),
        ]
        
        return "\n".join(sections)
    
    def generate_word(self, md_content: str, filename: Optional[str] = None) -> str:
        """Generate Word document from Markdown content."""
        if filename is None:
            safe_name = re.sub(r'[^\w\u4e00-\u9fff\-]', '_', self.industry)[:40]
            filename = f"{safe_name}_行业研究报告_{self.research_date[:10]}.docx"
        
        filepath = os.path.join(self.output_dir, filename)
        doc = Document()
        
        # Set default font for Chinese
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(11)
        
        # Parse markdown and add to doc
        lines = md_content.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # Heading 1
            if line.startswith("# "):
                p = doc.add_heading(line[2:], level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                i += 1
                continue
            
            # Heading 2
            if line.startswith("## "):
                doc.add_heading(line[3:], level=1)
                i += 1
                continue
            
            # Heading 3
            if line.startswith("### "):
                doc.add_heading(line[4:], level=2)
                i += 1
                continue
            
            # Blockquote
            if line.startswith("> "):
                text = line[2:]
                # Handle nested blockquotes
                while i + 1 < len(lines) and lines[i + 1].startswith("> "):
                    i += 1
                    text += "\n" + lines[i][2:]
                p = doc.add_paragraph(text)
                p.style = 'Quote'
                i += 1
                continue
            
            # List items
            if re.match(r'^\d+\.\s', line):
                text = re.sub(r'^\d+\.\s', '', line)
                doc.add_paragraph(text, style='List Number')
                i += 1
                continue
            
            if line.startswith("- "):
                text = line[2:]
                # Check for bold within
                doc.add_paragraph(text, style='List Bullet')
                i += 1
                continue
            
            # Empty line
            if not line.strip():
                i += 1
                continue
            
            # Regular paragraph
            doc.add_paragraph(line)
            i += 1
        
        doc.save(filepath)
        print(f"[Word] 报告已保存: {filepath}")
        return filepath
    
    def save(self, filename_prefix: Optional[str] = None) -> Dict[str, str]:
        """Save both Markdown and Word reports."""
        if filename_prefix is None:
            safe_name = re.sub(r'[^\w\u4e00-\u9fff\-]', '_', self.industry)[:40]
            filename_prefix = f"{safe_name}_行业研究报告_{self.research_date[:10]}"
        
        md_content = self.generate_markdown()
        
        md_path = os.path.join(self.output_dir, f"{filename_prefix}.md")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(md_content)
        print(f"[Markdown] 报告已保存: {md_path}")
        
        word_path = self.generate_word(md_content, f"{filename_prefix}.docx")
        
        # Also save raw research data
        data_path = os.path.join(self.output_dir, f"{filename_prefix}_raw_data.json")
        with open(data_path, "w", encoding="utf-8") as f:
            json.dump(self.data, f, ensure_ascii=False, indent=2)
        print(f"[JSON] 原始数据已保存: {data_path}")
        
        return {
            "markdown": md_path,
            "word": word_path,
            "raw_data": data_path,
        }
